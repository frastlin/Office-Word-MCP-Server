"""
Core comment extraction functionality for Word documents.

This module reads ``word/comments.xml`` directly from the .docx zip rather
than relying on python-docx's relationship graph. The relationship approach
was unreliable when external editors (docx-editor) added the comments part
or when the OPC package contained extra zip entries (e.g. the historical
docx-editor ``meta.json`` leak), causing ``get_all_comments`` to silently
return zero even when comments were visible in Word.
"""
import datetime
import zipfile
from typing import Dict, List, Optional, Any, Union

from docx import Document
from docx.document import Document as DocumentType
from docx.text.paragraph import Paragraph
from lxml import etree


W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
W15_NS = "http://schemas.microsoft.com/office/word/2012/wordml"
COMMENTS_PART = "word/comments.xml"


def extract_all_comments(doc_or_path: Union[DocumentType, str]) -> List[Dict[str, Any]]:
    """Extract all comments from a Word document.

    Accepts either a filename (preferred — reads comments.xml directly from
    the zip) or a python-docx ``Document`` object (resolved to a filename
    via the document part's ``package.filename`` when possible).
    """
    filename = _resolve_filename(doc_or_path)
    if filename is None:
        return []
    return _extract_from_zip(filename)


def _resolve_filename(doc_or_path: Union[DocumentType, str, None]) -> Optional[str]:
    if isinstance(doc_or_path, str):
        return doc_or_path
    if doc_or_path is None:
        return None
    # python-docx Document — try to recover the underlying file path.
    try:
        pkg = getattr(doc_or_path.part, "package", None)
        for attr in ("filename", "_filename", "partname"):
            value = getattr(pkg, attr, None)
            if isinstance(value, str):
                return value
    except Exception:
        pass
    return None


def _extract_from_zip(filename: str) -> List[Dict[str, Any]]:
    try:
        with zipfile.ZipFile(filename) as z:
            if COMMENTS_PART not in z.namelist():
                return []
            xml_bytes = z.read(COMMENTS_PART)
    except (zipfile.BadZipFile, FileNotFoundError, KeyError):
        return []

    try:
        root = etree.fromstring(xml_bytes)
    except etree.XMLSyntaxError:
        return []

    comments: List[Dict[str, Any]] = []
    for idx, elem in enumerate(root.findall(f"{{{W_NS}}}comment")):
        comment_data = extract_comment_data(elem, idx)
        if comment_data:
            comments.append(comment_data)
    return comments


def extract_comments_from_paragraphs(doc: DocumentType) -> List[Dict[str, Any]]:
    """
    Extract comments by scanning paragraphs for comment references.
    
    Args:
        doc: The Document object
        
    Returns:
        List of comment dictionaries
    """
    comments = []
    comment_id = 1
    
    # Check all paragraphs in the document
    for para_idx, paragraph in enumerate(doc.paragraphs):
        para_comments = find_paragraph_comments(paragraph, para_idx, comment_id)
        comments.extend(para_comments)
        comment_id += len(para_comments)
    
    # Check paragraphs in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para_idx, paragraph in enumerate(cell.paragraphs):
                    para_comments = find_paragraph_comments(paragraph, para_idx, comment_id, in_table=True)
                    comments.extend(para_comments)
                    comment_id += len(para_comments)
    
    return comments


def extract_comment_data(comment_element, index: int) -> Optional[Dict[str, Any]]:
    """
    Extract data from a comment XML element.
    
    Args:
        comment_element: The XML comment element
        index: Index for generating a unique ID
        
    Returns:
        Dictionary with comment data or None
    """
    try:
        # Extract comment attributes
        comment_id = comment_element.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}id', str(index))
        author = comment_element.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}author', 'Unknown')
        initials = comment_element.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}initials', '')
        date_str = comment_element.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}date', '')
        
        # Parse date if available
        date = None
        if date_str:
            try:
                date = datetime.datetime.fromisoformat(date_str.replace('Z', '+00:00'))
                date = date.isoformat()
            except:
                date = date_str
        
        # Extract comment text (handle lxml Elements rooted from a parsed
        # comments.xml as well as python-docx's xpath-aware elements)
        text_elements = comment_element.xpath(
            './/w:t', namespaces={'w': W_NS}
        )
        text = ''.join(elem.text or '' for elem in text_elements)
        
        return {
            'id': f'comment_{index + 1}',
            'comment_id': comment_id,
            'author': author,
            'initials': initials,
            'date': date,
            'text': text.strip(),
            'paragraph_index': None,  # Will be filled if we can determine it
            'in_table': False,
            'reference_text': ''
        }
    
    except Exception as e:
        return None


def find_paragraph_comments(paragraph: Paragraph, para_index: int, 
                           start_id: int, in_table: bool = False) -> List[Dict[str, Any]]:
    """
    Find comments associated with a specific paragraph.
    
    Args:
        paragraph: The paragraph to check
        para_index: The index of the paragraph
        start_id: Starting ID for comments
        in_table: Whether the paragraph is in a table
        
    Returns:
        List of comment dictionaries
    """
    comments = []
    
    try:
        # Access the paragraph's XML element
        para_xml = paragraph._element
        
        # Look for comment range markers (simplified approach)
        # This is a basic implementation - the full version would need more sophisticated XML parsing
        xml_text = str(para_xml)
        
        # Simple check for comment references in the XML
        if 'commentRangeStart' in xml_text or 'commentReference' in xml_text:
            # Create a placeholder comment entry
            comment_info = {
                'id': f'comment_{start_id}',
                'comment_id': f'{start_id}',
                'author': 'Unknown',
                'initials': '',
                'date': None,
                'text': 'Comment detected but content not accessible',
                'paragraph_index': para_index,
                'in_table': in_table,
                'reference_text': paragraph.text[:50] + '...' if len(paragraph.text) > 50 else paragraph.text
            }
            comments.append(comment_info)
    
    except Exception:
        # If we can't access the XML, skip this paragraph
        pass
    
    return comments


def filter_comments_by_author(comments: List[Dict[str, Any]], author: str) -> List[Dict[str, Any]]:
    """
    Filter comments by author name.
    
    Args:
        comments: List of comment dictionaries
        author: Author name to filter by (case-insensitive)
        
    Returns:
        Filtered list of comments
    """
    author_lower = author.lower()
    return [c for c in comments if c.get('author', '').lower() == author_lower]


def get_comments_for_paragraph(comments: List[Dict[str, Any]], paragraph_index: int) -> List[Dict[str, Any]]:
    """
    Get all comments for a specific paragraph.
    
    Args:
        comments: List of all comments
        paragraph_index: Index of the paragraph
        
    Returns:
        Comments for the specified paragraph
    """
    return [c for c in comments if c.get('paragraph_index') == paragraph_index]