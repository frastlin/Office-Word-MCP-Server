"""Tests for replace_paragraph_text with markdown formatting support."""
import pytest
from docx import Document
from word_document_server.utils.document_utils import replace_paragraph_text, _parse_markdown_runs


class TestParseMarkdownRuns:
    """Unit tests for the markdown parser."""

    def test_plain_text(self):
        """No markdown returns single plain run."""
        runs = _parse_markdown_runs("Hello world")
        assert len(runs) == 1
        assert runs[0]["text"] == "Hello world"
        assert runs[0]["bold"] is False
        assert runs[0]["italic"] is False

    def test_single_italic(self):
        """'Hello *world*' returns 2 runs, second is italic."""
        runs = _parse_markdown_runs("Hello *world*")
        assert len(runs) == 2
        assert runs[0]["text"] == "Hello "
        assert runs[0]["italic"] is False
        assert runs[1]["text"] == "world"
        assert runs[1]["italic"] is True
        assert runs[1]["bold"] is False

    def test_single_bold(self):
        """'Hello **world**' returns 2 runs, second is bold."""
        runs = _parse_markdown_runs("Hello **world**")
        assert len(runs) == 2
        assert runs[1]["text"] == "world"
        assert runs[1]["bold"] is True
        assert runs[1]["italic"] is False

    def test_bold_italic(self):
        """'Hello ***world***' returns 2 runs, second is bold+italic."""
        runs = _parse_markdown_runs("Hello ***world***")
        assert len(runs) == 2
        assert runs[1]["text"] == "world"
        assert runs[1]["bold"] is True
        assert runs[1]["italic"] is True

    def test_mixed_formatting(self):
        """'The *F* statistic (*p* < .001)' produces correct runs."""
        runs = _parse_markdown_runs("The *F* statistic (*p* < .001)")
        # Expected: "The " + italic "F" + " statistic (" + italic "p" + " < .001)"
        texts = [r["text"] for r in runs]
        assert "".join(texts) == "The F statistic (p < .001)"
        # Find italic runs
        italic_runs = [r for r in runs if r["italic"]]
        italic_texts = [r["text"] for r in italic_runs]
        assert "F" in italic_texts
        assert "p" in italic_texts

    def test_adjacent_bold_and_italic(self):
        """'**bold** and *italic*' produces 3 runs."""
        runs = _parse_markdown_runs("**bold** and *italic*")
        texts = [r["text"] for r in runs]
        assert "".join(texts) == "bold and italic"
        bold_runs = [r for r in runs if r["bold"]]
        assert any(r["text"] == "bold" for r in bold_runs)
        italic_runs = [r for r in runs if r["italic"]]
        assert any(r["text"] == "italic" for r in italic_runs)

    def test_empty_string(self):
        """Empty string returns single empty run."""
        runs = _parse_markdown_runs("")
        assert len(runs) == 1
        assert runs[0]["text"] == ""

    def test_no_closing_marker(self):
        """Unclosed marker is treated as literal text."""
        runs = _parse_markdown_runs("Hello *world")
        full_text = "".join(r["text"] for r in runs)
        assert full_text == "Hello *world"

    def test_asterisks_in_math(self):
        """Single asterisks not used as markers when not paired."""
        runs = _parse_markdown_runs("2 * 3 = 6")
        full_text = "".join(r["text"] for r in runs)
        assert full_text == "2 * 3 = 6"


class TestReplaceWithMarkdown:
    """Integration tests: markdown formatting applied to Word document."""

    def test_italic_run_in_document(self, make_docx):
        """parse_markdown=True creates italic run in saved document."""
        path = make_docx(paragraphs=["Original text"])
        replace_paragraph_text(path, 0, "Hello *world*", parse_markdown=True)
        doc = Document(path)
        para = doc.paragraphs[0]
        assert para.text == "Hello world"
        # Find the italic run
        italic_runs = [r for r in para.runs if r.italic]
        assert len(italic_runs) >= 1
        assert any(r.text == "world" for r in italic_runs)

    def test_bold_run_in_document(self, make_docx):
        """parse_markdown=True creates bold run in saved document."""
        path = make_docx(paragraphs=["Original text"])
        replace_paragraph_text(path, 0, "Hello **world**", parse_markdown=True)
        doc = Document(path)
        para = doc.paragraphs[0]
        bold_runs = [r for r in para.runs if r.bold]
        assert any(r.text == "world" for r in bold_runs)

    def test_bold_italic_in_document(self, make_docx):
        """parse_markdown=True creates bold+italic run."""
        path = make_docx(paragraphs=["Original"])
        replace_paragraph_text(path, 0, "Hello ***world***", parse_markdown=True)
        doc = Document(path)
        para = doc.paragraphs[0]
        bi_runs = [r for r in para.runs if r.bold and r.italic]
        assert any(r.text == "world" for r in bi_runs)

    def test_preserves_paragraph_style(self, make_docx):
        """Heading 2 style preserved with parse_markdown=True."""
        path = make_docx(paragraphs=[
            {"style": "Heading 2", "runs": [{"text": "Old heading"}]},
        ])
        replace_paragraph_text(path, 0, "New *heading*", preserve_style=True, parse_markdown=True)
        doc = Document(path)
        assert doc.paragraphs[0].style.name == "Heading 2"

    def test_default_no_parsing(self, make_docx):
        """parse_markdown=False (default) treats asterisks as literal text."""
        path = make_docx(paragraphs=["Original"])
        replace_paragraph_text(path, 0, "Hello *world*")
        doc = Document(path)
        assert doc.paragraphs[0].text == "Hello *world*"

    def test_statistical_text(self, make_docx):
        """Realistic: 'ANOVA (*F*(2, 38) = 108.37, *p* < .001)' formats correctly."""
        path = make_docx(paragraphs=["Old text"])
        new_text = "ANOVA (*F*(2, 38) = 108.37, *p* < .001)"
        replace_paragraph_text(path, 0, new_text, parse_markdown=True)
        doc = Document(path)
        para = doc.paragraphs[0]
        # Full text should have no asterisks
        assert para.text == "ANOVA (F(2, 38) = 108.37, p < .001)"
        # F and p should be italic
        italic_texts = [r.text for r in para.runs if r.italic]
        assert "F" in italic_texts
        assert "p" in italic_texts

    def test_multiple_paragraphs_unchanged(self, make_docx):
        """Other paragraphs in the document are not affected."""
        path = make_docx(paragraphs=["Keep me", "Replace me", "Keep me too"])
        replace_paragraph_text(path, 1, "New *italic* text", parse_markdown=True)
        doc = Document(path)
        assert doc.paragraphs[0].text == "Keep me"
        assert doc.paragraphs[2].text == "Keep me too"
