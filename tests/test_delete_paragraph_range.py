"""Tests for delete_paragraph_range utility."""
import pytest
from docx import Document
from word_document_server.utils.document_utils import delete_paragraph_range


class TestDeleteRange:
    """Core deletion behavior."""

    def test_delete_middle_range(self, make_docx):
        """Delete paragraphs 1-3 from 5-paragraph doc, verify A and E remain."""
        path = make_docx(paragraphs=["A", "B", "C", "D", "E"])
        result = delete_paragraph_range(path, 1, 3)
        assert "error" not in result.lower()
        doc = Document(path)
        texts = [p.text for p in doc.paragraphs]
        assert texts == ["A", "E"]

    def test_delete_single_paragraph_range(self, make_docx):
        """start_index == end_index deletes exactly one paragraph."""
        path = make_docx(paragraphs=["A", "B", "C"])
        delete_paragraph_range(path, 1, 1)
        doc = Document(path)
        texts = [p.text for p in doc.paragraphs]
        assert texts == ["A", "C"]

    def test_delete_from_start(self, make_docx):
        """Delete paragraphs 0-2 from 5-paragraph doc."""
        path = make_docx(paragraphs=["A", "B", "C", "D", "E"])
        delete_paragraph_range(path, 0, 2)
        doc = Document(path)
        texts = [p.text for p in doc.paragraphs]
        assert texts == ["D", "E"]

    def test_delete_to_end(self, make_docx):
        """Delete last 3 paragraphs of 5-paragraph doc."""
        path = make_docx(paragraphs=["A", "B", "C", "D", "E"])
        delete_paragraph_range(path, 2, 4)
        doc = Document(path)
        texts = [p.text for p in doc.paragraphs]
        assert texts == ["A", "B"]

    def test_delete_all_paragraphs(self, make_docx):
        """Delete all paragraphs from doc."""
        path = make_docx(paragraphs=["A", "B", "C"])
        delete_paragraph_range(path, 0, 2)
        doc = Document(path)
        assert len(doc.paragraphs) == 0 or all(p.text == "" for p in doc.paragraphs)


class TestDeleteRangeValidation:
    """Input validation."""

    def test_start_greater_than_end(self, make_docx):
        """start_index > end_index returns error."""
        path = make_docx(paragraphs=["A", "B", "C"])
        result = delete_paragraph_range(path, 2, 1)
        assert "error" in result.lower()

    def test_end_out_of_bounds(self, make_docx):
        """end_index beyond doc length returns error."""
        path = make_docx(paragraphs=["A", "B", "C"])
        result = delete_paragraph_range(path, 0, 10)
        assert "error" in result.lower()

    def test_negative_start(self, make_docx):
        """Negative start_index returns error."""
        path = make_docx(paragraphs=["A", "B", "C"])
        result = delete_paragraph_range(path, -1, 2)
        assert "error" in result.lower()

    def test_file_not_found(self, tmp_path):
        """Non-existent file returns error."""
        result = delete_paragraph_range(str(tmp_path / "missing.docx"), 0, 1)
        assert "error" in result.lower() or "not exist" in result.lower()


class TestDeleteRangePreservation:
    """Verify surrounding content is untouched."""

    def test_surrounding_paragraphs_unchanged(self, make_docx):
        """Paragraphs before start and after end retain their text and style."""
        path = make_docx(paragraphs=[
            {"style": "Heading 1", "runs": [{"text": "Title"}]},
            "Delete me",
            "Delete me too",
            {"style": "Heading 2", "runs": [{"text": "Subtitle"}]},
        ])
        delete_paragraph_range(path, 1, 2)
        doc = Document(path)
        assert doc.paragraphs[0].text == "Title"
        assert doc.paragraphs[0].style.name == "Heading 1"
        assert doc.paragraphs[1].text == "Subtitle"
        assert doc.paragraphs[1].style.name == "Heading 2"

    def test_paragraph_count_reduced(self, make_docx):
        """Total paragraph count decreases by (end - start + 1)."""
        path = make_docx(paragraphs=["A", "B", "C", "D", "E"])
        delete_paragraph_range(path, 1, 3)
        doc = Document(path)
        assert len(doc.paragraphs) == 2
