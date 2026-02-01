"""Tests for replace_paragraph_text tool (Shortcoming 1)."""
import pytest
from docx import Document

from word_document_server.utils.document_utils import replace_paragraph_text


class TestBasicReplacement:
    def test_replace_middle_paragraph(self, make_docx):
        path = make_docx(paragraphs=["First", "Second", "Third"])
        result = replace_paragraph_text(path, 1, "Replaced")
        assert "successfully" in result.lower()
        doc = Document(path)
        assert doc.paragraphs[0].text == "First"
        assert doc.paragraphs[1].text == "Replaced"
        assert doc.paragraphs[2].text == "Third"

    def test_paragraph_count_unchanged(self, make_docx):
        path = make_docx(paragraphs=["First", "Second", "Third"])
        doc_before = Document(path)
        count_before = len(doc_before.paragraphs)
        replace_paragraph_text(path, 1, "Replaced")
        doc_after = Document(path)
        assert len(doc_after.paragraphs) == count_before


class TestStylePreservation:
    def test_preserves_style_by_default(self, make_docx):
        path = make_docx(paragraphs=[
            {"style": "Heading 2", "runs": [{"text": "My Heading"}]},
            "Body text",
        ])
        replace_paragraph_text(path, 0, "New Heading Text")
        doc = Document(path)
        assert doc.paragraphs[0].text == "New Heading Text"
        assert doc.paragraphs[0].style.name == "Heading 2"

    def test_preserve_style_false(self, make_docx):
        path = make_docx(paragraphs=[
            {"style": "Heading 2", "runs": [{"text": "My Heading"}]},
            "Body text",
        ])
        replace_paragraph_text(path, 0, "New Text", preserve_style=False)
        doc = Document(path)
        assert doc.paragraphs[0].text == "New Text"
        assert doc.paragraphs[0].style.name == "Normal"


class TestInvalidIndex:
    def test_index_too_large(self, make_docx):
        path = make_docx(paragraphs=["Only one"])
        result = replace_paragraph_text(path, 5, "New text")
        assert "invalid" in result.lower()

    def test_negative_index(self, make_docx):
        path = make_docx(paragraphs=["Only one"])
        result = replace_paragraph_text(path, -1, "New text")
        assert "invalid" in result.lower()
