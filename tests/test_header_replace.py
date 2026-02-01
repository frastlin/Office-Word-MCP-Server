"""Tests for header matching normalization (Bug 3)."""
import pytest
from docx import Document

from word_document_server.utils.document_utils import replace_paragraph_block_below_header


class TestExactHeadingMatch:
    """Regression: exact heading match still works."""

    def test_exact_heading_replaces_content(self, heading_docx):
        result = replace_paragraph_block_below_header(
            heading_docx,
            header_text="Section One",
            new_paragraphs=["New content A", "New content B"],
        )
        assert "not found" not in result.lower()
        doc = Document(heading_docx)
        texts = [p.text for p in doc.paragraphs]
        assert "New content A" in texts
        assert "New content B" in texts
        assert "Content under section one." not in texts
        assert "More content." not in texts
        # Section Two should be untouched
        assert "Content under section two." in texts


class TestHeadingNBSP:
    """NBSP (U+00A0) in heading should match regular space in search text."""

    def test_nbsp_heading_matches(self, tmp_path):
        path = tmp_path / "nbsp_heading.docx"
        doc = Document()
        p = doc.add_paragraph("Section\u00a0One", style="Heading 1")
        doc.add_paragraph("Content under section one.")
        doc.add_paragraph("More content.")
        p2 = doc.add_paragraph("Section Two", style="Heading 1")
        doc.add_paragraph("Content under section two.")
        doc.save(str(path))

        result = replace_paragraph_block_below_header(
            str(path),
            header_text="Section One",
            new_paragraphs=["Replaced"],
        )
        assert "not found" not in result.lower()
        doc2 = Document(str(path))
        texts = [p.text for p in doc2.paragraphs]
        assert "Replaced" in texts
        assert "Content under section one." not in texts


class TestHeadingAutoNumbering:
    """Auto-numbering prefix should match via contains + heading style check."""

    def test_numbered_heading_matches(self, tmp_path):
        path = tmp_path / "numbered_heading.docx"
        doc = Document()
        p = doc.add_paragraph("1. Section One", style="Heading 1")
        doc.add_paragraph("Content under section one.")
        p2 = doc.add_paragraph("2. Section Two", style="Heading 1")
        doc.add_paragraph("Content under section two.")
        doc.save(str(path))

        result = replace_paragraph_block_below_header(
            str(path),
            header_text="Section One",
            new_paragraphs=["Replaced"],
        )
        assert "not found" not in result.lower()


class TestHeadingNotFound:
    """Non-existent heading text should return error."""

    def test_heading_not_found(self, heading_docx):
        result = replace_paragraph_block_below_header(
            heading_docx,
            header_text="NONEXISTENT",
            new_paragraphs=["New stuff"],
        )
        assert "not found" in result.lower()


class TestStopsAtNextHeading:
    """Replacement should stop at the next heading."""

    def test_stops_at_next_heading(self, heading_docx):
        result = replace_paragraph_block_below_header(
            heading_docx,
            header_text="Section One",
            new_paragraphs=["Only this"],
        )
        doc = Document(heading_docx)
        texts = [p.text for p in doc.paragraphs]
        # Section Two header and content should still be there
        assert "Section Two" in texts
        assert "Content under section two." in texts
