"""Tests for anchor matching normalization (Bug 2)."""
import pytest
from docx import Document

from word_document_server.utils.document_utils import replace_block_between_manual_anchors


class TestAnchorExactMatch:
    """Regression: exact anchor matching still works."""

    def test_exact_match_replaces_content(self, anchor_docx):
        result = replace_block_between_manual_anchors(
            anchor_docx,
            start_anchor_text="--- START ANCHOR ---",
            end_anchor_text="--- END ANCHOR ---",
            new_paragraphs=["New content A", "New content B"],
        )
        assert "not found" not in result.lower()
        doc = Document(anchor_docx)
        texts = [p.text for p in doc.paragraphs]
        assert "New content A" in texts
        assert "New content B" in texts
        assert "Content to replace 1" not in texts
        assert "Content to replace 2" not in texts


class TestAnchorNBSP:
    """NBSP (U+00A0) in document should match regular space in anchor text."""

    def test_nbsp_anchor_matches(self, nbsp_anchor_docx):
        result = replace_block_between_manual_anchors(
            nbsp_anchor_docx,
            start_anchor_text="--- START ANCHOR ---",
            end_anchor_text="--- END ANCHOR ---",
            new_paragraphs=["Replaced content"],
        )
        assert "not found" not in result.lower()
        doc = Document(nbsp_anchor_docx)
        texts = [p.text for p in doc.paragraphs]
        assert "Replaced content" in texts
        assert "Content to replace" not in texts


class TestAnchorExtraWhitespace:
    """Extra leading/trailing whitespace should be handled."""

    def test_extra_whitespace_matches(self, tmp_path):
        path = tmp_path / "ws_test.docx"
        doc = Document()
        doc.add_paragraph("  --- START ANCHOR ---  ")
        doc.add_paragraph("Content to replace")
        doc.add_paragraph("  --- END ANCHOR ---  ")
        doc.save(str(path))

        result = replace_block_between_manual_anchors(
            str(path),
            start_anchor_text="--- START ANCHOR ---",
            end_anchor_text="--- END ANCHOR ---",
            new_paragraphs=["New stuff"],
        )
        assert "not found" not in result.lower()


class TestAnchorContainsFallback:
    """Auto-numbering prefix like '1. ' should match via contains fallback."""

    def test_numbered_anchor_matches(self, tmp_path):
        path = tmp_path / "numbered_test.docx"
        doc = Document()
        doc.add_paragraph("1. --- START ANCHOR ---")
        doc.add_paragraph("Content to replace")
        doc.add_paragraph("2. --- END ANCHOR ---")
        doc.save(str(path))

        result = replace_block_between_manual_anchors(
            str(path),
            start_anchor_text="--- START ANCHOR ---",
            end_anchor_text="--- END ANCHOR ---",
            new_paragraphs=["New stuff"],
        )
        assert "not found" not in result.lower()


class TestAnchorNotFound:
    """Non-existent anchor text should return error."""

    def test_anchor_not_found(self, anchor_docx):
        result = replace_block_between_manual_anchors(
            anchor_docx,
            start_anchor_text="NONEXISTENT",
            end_anchor_text="ALSO NONEXISTENT",
            new_paragraphs=["New stuff"],
        )
        assert "not found" in result.lower()
