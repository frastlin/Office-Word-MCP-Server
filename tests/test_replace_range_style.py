"""Tests for replace_paragraph_range style behavior."""
import pytest
from docx import Document
from word_document_server.utils.document_utils import replace_paragraph_range


class TestReplaceRangeDefaultStyle:
    """Default style behavior when no style/preserve_style specified."""

    def test_default_style_is_normal(self, make_docx):
        """No style or preserve_style: new paragraphs get Normal style."""
        path = make_docx(paragraphs=[
            {"style": "Heading 2", "runs": [{"text": "Old heading"}]},
            "Old content",
        ])
        replace_paragraph_range(path, 0, 1, ["New text"])
        doc = Document(path)
        assert doc.paragraphs[0].style.name == "Normal"


class TestReplaceRangePreserveStyle:
    """preserve_style=True behavior."""

    def test_preserve_heading_style(self, make_docx):
        """Replace Heading 2 paragraphs with preserve_style=True, new paras get Heading 2."""
        path = make_docx(paragraphs=[
            {"style": "Heading 2", "runs": [{"text": "Old heading"}]},
            "Old content",
            "Another paragraph",
        ])
        replace_paragraph_range(path, 0, 0, ["New heading text"], preserve_style=True)
        doc = Document(path)
        assert doc.paragraphs[0].style.name == "Heading 2"
        assert doc.paragraphs[0].text == "New heading text"

    def test_style_param_overrides_preserve(self, make_docx):
        """Explicit style param takes precedence over preserve_style."""
        path = make_docx(paragraphs=[
            {"style": "Heading 1", "runs": [{"text": "Title"}]},
            "Content",
        ])
        replace_paragraph_range(path, 0, 0, ["New text"],
                                style="Heading 3", preserve_style=True)
        doc = Document(path)
        assert doc.paragraphs[0].style.name == "Heading 3"

    def test_preserve_style_multiple_paragraphs(self, make_docx):
        """All new paragraphs get the preserved style."""
        path = make_docx(paragraphs=[
            {"style": "Heading 2", "runs": [{"text": "Old"}]},
            "Content after",
        ])
        replace_paragraph_range(path, 0, 0, ["New A", "New B"], preserve_style=True)
        doc = Document(path)
        assert doc.paragraphs[0].style.name == "Heading 2"
        assert doc.paragraphs[1].style.name == "Heading 2"

    def test_preserve_false_same_as_default(self, make_docx):
        """preserve_style=False behaves same as default (Normal)."""
        path = make_docx(paragraphs=[
            {"style": "Heading 1", "runs": [{"text": "Title"}]},
        ])
        replace_paragraph_range(path, 0, 0, ["New text"], preserve_style=False)
        doc = Document(path)
        assert doc.paragraphs[0].style.name == "Normal"


class TestReplaceRangeEmptySpacers:
    """Empty string handling for spacer paragraphs (Issue 8)."""

    def test_empty_string_creates_spacer(self, make_docx):
        """Empty string in new_paragraphs creates empty paragraph."""
        path = make_docx(paragraphs=["Old content"])
        replace_paragraph_range(path, 0, 0, ["Content", "", "More content"])
        doc = Document(path)
        texts = [p.text for p in doc.paragraphs]
        assert texts == ["Content", "", "More content"]

    def test_spacers_get_style(self, make_docx):
        """Empty spacer paragraphs also get the specified/preserved style."""
        path = make_docx(paragraphs=[
            {"style": "Heading 2", "runs": [{"text": "Title"}]},
        ])
        replace_paragraph_range(path, 0, 0, ["Text", ""], preserve_style=True)
        doc = Document(path)
        # Even the empty spacer should have the preserved style
        assert doc.paragraphs[1].style.name == "Heading 2"


class TestReplaceRangeEmptyList:
    """Empty new_paragraphs list behavior."""

    def test_empty_list_deletes_range(self, make_docx):
        """new_paragraphs=[] effectively deletes the range."""
        path = make_docx(paragraphs=["A", "B", "C", "D"])
        replace_paragraph_range(path, 1, 2, [])
        doc = Document(path)
        texts = [p.text for p in doc.paragraphs]
        assert texts == ["A", "D"]
