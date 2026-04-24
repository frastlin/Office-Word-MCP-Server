"""Tests for hyperlink support (add, mixed segments, convert, markdown)."""
import asyncio
import zipfile
from xml.etree import ElementTree as ET

import pytest
from docx import Document

from word_document_server.tools import content_tools


W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
R_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
REL_PKG_NS = "http://schemas.openxmlformats.org/package/2006/relationships"


def _run(coro):
    return asyncio.get_event_loop().run_until_complete(coro) if False else asyncio.run(coro)


def _iter_hyperlink_elements(path):
    with zipfile.ZipFile(path) as z:
        doc_xml = z.read("word/document.xml")
    root = ET.fromstring(doc_xml)
    return root.findall(f".//{{{W_NS}}}hyperlink")


def _external_rels(path):
    with zipfile.ZipFile(path) as z:
        rels_xml = z.read("word/_rels/document.xml.rels")
    root = ET.fromstring(rels_xml)
    hyper_type = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink"
    return [
        r for r in root.findall(f"{{{REL_PKG_NS}}}Relationship")
        if r.get("Type") == hyper_type
    ]


def test_add_hyperlink_creates_element_and_external_rel(make_docx):
    path = make_docx(paragraphs=["Intro"])
    result = _run(content_tools.add_hyperlink(
        path,
        "https://xrnavigation.io/audiom-demo",
        "Interactive Audiom Demo",
    ))
    assert "Hyperlink added" in result

    links = _iter_hyperlink_elements(path)
    assert len(links) == 1
    r_id = links[0].get(f"{{{R_NS}}}id")
    assert r_id

    rels = _external_rels(path)
    assert len(rels) == 1
    rel = rels[0]
    assert rel.get("Id") == r_id
    assert rel.get("Target") == "https://xrnavigation.io/audiom-demo"
    assert rel.get("TargetMode") == "External"

    # Visible text lives in the w:t inside the hyperlink.
    t_elem = links[0].find(f".//{{{W_NS}}}t")
    assert t_elem is not None
    assert t_elem.text == "Interactive Audiom Demo"


def test_add_hyperlink_defaults_text_to_url(make_docx):
    path = make_docx(paragraphs=[])
    _run(content_tools.add_hyperlink(path, "https://example.com"))
    links = _iter_hyperlink_elements(path)
    t = links[0].find(f".//{{{W_NS}}}t")
    assert t.text == "https://example.com"


def test_add_hyperlink_adds_missing_scheme(make_docx):
    path = make_docx(paragraphs=[])
    _run(content_tools.add_hyperlink(path, "example.com", "Example"))
    rels = _external_rels(path)
    assert rels[0].get("Target") == "https://example.com"


def test_add_hyperlink_dedup_same_url(make_docx):
    path = make_docx(paragraphs=[])
    _run(content_tools.add_hyperlink(path, "https://example.com", "A"))
    _run(content_tools.add_hyperlink(path, "https://example.com", "B"))

    links = _iter_hyperlink_elements(path)
    assert len(links) == 2
    rels = _external_rels(path)
    assert len(rels) == 1  # single relationship for shared URL


def test_add_paragraph_with_hyperlinks_mixed(make_docx):
    path = make_docx(paragraphs=[])
    segments = [
        {"text": "See the "},
        {"text": "Interactive Audiom Demo", "url": "https://xrnavigation.io/audiom-demo"},
        {"text": " for details."},
    ]
    _run(content_tools.add_paragraph_with_hyperlinks(path, segments))

    doc = Document(path)
    last_para = doc.paragraphs[-1]
    # Combined text across plain runs and hyperlink runs should match.
    assert last_para.text == "See the Interactive Audiom Demo for details."

    links = _iter_hyperlink_elements(path)
    assert len(links) == 1
    t = links[0].find(f".//{{{W_NS}}}t")
    assert t.text == "Interactive Audiom Demo"


def test_convert_markdown_links(make_docx):
    path = make_docx(paragraphs=[
        "Visit [Interactive Audiom Demo](https://xrnavigation.io/audiom-demo) today.",
        "No link here.",
    ])
    result = _run(content_tools.convert_markdown_links(path))
    assert "Converted 1" in result

    doc = Document(path)
    assert doc.paragraphs[0].text == "Visit Interactive Audiom Demo today."

    links = _iter_hyperlink_elements(path)
    assert len(links) == 1
    t = links[0].find(f".//{{{W_NS}}}t")
    assert t.text == "Interactive Audiom Demo"

    rels = _external_rels(path)
    assert rels[0].get("Target") == "https://xrnavigation.io/audiom-demo"


def test_convert_markdown_links_multiple_in_one_paragraph(make_docx):
    path = make_docx(paragraphs=[
        "See [A](https://a.example.com) and [B](https://b.example.com).",
    ])
    _run(content_tools.convert_markdown_links(path))
    doc = Document(path)
    assert doc.paragraphs[0].text == "See A and B."
    assert len(_iter_hyperlink_elements(path)) == 2
    targets = {r.get("Target") for r in _external_rels(path)}
    assert targets == {"https://a.example.com", "https://b.example.com"}


def test_convert_markdown_links_none_found(make_docx):
    path = make_docx(paragraphs=["Just plain text."])
    result = _run(content_tools.convert_markdown_links(path))
    assert "No markdown links" in result


def test_convert_text_to_hyperlink_basic(make_docx):
    path = make_docx(paragraphs=[
        "Visit Interactive Audiom Demo today.",
    ])
    result = _run(content_tools.convert_text_to_hyperlink(
        path,
        "Interactive Audiom Demo",
        "https://xrnavigation.io/audiom-demo",
    ))
    assert "Converted 1" in result

    doc = Document(path)
    assert doc.paragraphs[0].text == "Visit Interactive Audiom Demo today."
    links = _iter_hyperlink_elements(path)
    assert len(links) == 1
    assert links[0].find(f".//{{{W_NS}}}t").text == "Interactive Audiom Demo"


def test_convert_text_to_hyperlink_cross_run(cross_run_docx):
    # 'Hello World' is split across two runs in the first paragraph.
    result = _run(content_tools.convert_text_to_hyperlink(
        cross_run_docx,
        "Hello World",
        "https://example.com",
    ))
    assert "Converted 1" in result

    doc = Document(cross_run_docx)
    assert doc.paragraphs[0].text == "Hello World"
    links = _iter_hyperlink_elements(cross_run_docx)
    assert len(links) == 1


def test_convert_text_to_hyperlink_not_found(make_docx):
    path = make_docx(paragraphs=["Nothing to see here."])
    result = _run(content_tools.convert_text_to_hyperlink(
        path, "Missing", "https://example.com",
    ))
    assert "not found" in result


def test_convert_text_to_hyperlink_every_occurrence(make_docx):
    path = make_docx(paragraphs=[
        "Click Link and then Link again.",
    ])
    _run(content_tools.convert_text_to_hyperlink(
        path, "Link", "https://example.com", occurrence=0,
    ))
    assert len(_iter_hyperlink_elements(path)) == 2


def test_file_reopens_after_operations(make_docx):
    path = make_docx(paragraphs=[])
    _run(content_tools.add_hyperlink(path, "https://example.com", "Example"))
    _run(content_tools.add_paragraph_with_hyperlinks(path, [
        {"text": "Plain "},
        {"text": "Clicky", "url": "https://example.org"},
    ]))
    _run(content_tools.convert_markdown_links(path))
    # Re-open without error.
    doc = Document(path)
    assert len(doc.paragraphs) >= 2
