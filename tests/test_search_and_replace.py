"""Tests for search_and_replace cross-run text matching (Bug 1)."""
import pytest
from docx import Document

from word_document_server.utils.document_utils import find_and_replace_text


class TestSingleRunReplace:
    """Regression: replacement within a single run still works."""

    def test_single_run_replace(self, make_docx):
        path = make_docx(paragraphs=["Hello World"])
        doc = Document(path)
        count = find_and_replace_text(doc, "Hello", "Goodbye")
        doc.save(path)
        doc2 = Document(path)
        assert doc2.paragraphs[0].text == "Goodbye World"
        assert count >= 1


class TestCrossRunReplace:
    """Text that spans multiple <w:r> elements must be matched and replaced."""

    def test_cross_run_replace(self, cross_run_docx):
        doc = Document(cross_run_docx)
        count = find_and_replace_text(doc, "Hello World", "Goodbye Earth")
        doc.save(cross_run_docx)
        doc2 = Document(cross_run_docx)
        assert doc2.paragraphs[0].text == "Goodbye Earth"
        assert count >= 1

    def test_preserves_first_run_formatting(self, multi_run_formatted_docx):
        doc = Document(multi_run_formatted_docx)
        count = find_and_replace_text(doc, "Hello World", "Goodbye Earth")
        doc.save(multi_run_formatted_docx)
        doc2 = Document(multi_run_formatted_docx)
        assert doc2.paragraphs[0].text == "Goodbye Earth"
        first_run = doc2.paragraphs[0].runs[0]
        assert first_run.bold is True
        assert count >= 1

    def test_table_cell_cross_run(self, table_docx):
        doc = Document(table_docx)
        count = find_and_replace_text(doc, "Hello World", "Goodbye Earth")
        doc.save(table_docx)
        doc2 = Document(table_docx)
        cell_text = doc2.tables[0].cell(0, 0).text
        assert cell_text == "Goodbye Earth"
        assert count >= 1

    def test_no_match_returns_zero(self, make_docx):
        path = make_docx(paragraphs=["Hello World"])
        doc = Document(path)
        count = find_and_replace_text(doc, "Nonexistent", "Replacement")
        assert count == 0

    def test_multiple_occurrences(self, make_docx):
        path = make_docx(paragraphs=[
            {"runs": [{"text": "Hello "}, {"text": "World"}]},
            {"runs": [{"text": "Hello "}, {"text": "World"}]},
        ])
        doc = Document(path)
        count = find_and_replace_text(doc, "Hello World", "Goodbye Earth")
        doc.save(path)
        doc2 = Document(path)
        assert doc2.paragraphs[0].text == "Goodbye Earth"
        assert doc2.paragraphs[1].text == "Goodbye Earth"
        assert count == 2

    def test_toc_paragraphs_skipped(self, tmp_path):
        path = tmp_path / "toc_test.docx"
        doc = Document()
        p = doc.add_paragraph("Hello World")
        p.style = doc.styles.add_style("TOC 1", 1) if "TOC 1" not in [s.name for s in doc.styles] else doc.styles["TOC 1"]
        doc.save(str(path))
        doc2 = Document(str(path))
        count = find_and_replace_text(doc2, "Hello World", "Goodbye Earth")
        assert count == 0
        assert doc2.paragraphs[0].text == "Hello World"
