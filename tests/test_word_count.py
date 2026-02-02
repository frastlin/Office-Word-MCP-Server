"""Tests for word count reporting in get_document_properties."""
import pytest
from docx import Document
from word_document_server.utils.document_utils import get_document_properties


class TestWordCountBody:
    """Body paragraph word counting."""

    def test_word_count_matches_split(self, make_docx):
        """Word count matches whitespace-split of body paragraphs."""
        path = make_docx(paragraphs=["Hello world", "One two three"])
        result = get_document_properties(path)
        assert result["word_count"] == 5  # 2 + 3

    def test_empty_paragraphs_zero_words(self, make_docx):
        """Empty paragraphs contribute 0 to word count."""
        path = make_docx(paragraphs=["Hello", "", "World", ""])
        result = get_document_properties(path)
        assert result["word_count"] == 2


class TestWordCountMetadata:
    """Word count metadata fields."""

    def test_word_count_method_present(self, make_docx):
        """Response includes word_count_method field."""
        path = make_docx(paragraphs=["Content"])
        result = get_document_properties(path)
        assert "word_count_method" in result
        assert isinstance(result["word_count_method"], str)

    def test_word_count_note_present(self, make_docx):
        """Response includes word_count_note field."""
        path = make_docx(paragraphs=["Content"])
        result = get_document_properties(path)
        assert "word_count_note" in result
        assert "table" in result["word_count_note"].lower() or "footnote" in result["word_count_note"].lower()


class TestTableWordCount:
    """Separate table word count."""

    def test_table_word_count_present(self, tmp_path):
        """Response includes table_word_count field."""
        path = tmp_path / "table_test.docx"
        doc = Document()
        doc.add_paragraph("Body text here")
        table = doc.add_table(rows=1, cols=1)
        table.cell(0, 0).text = "Table words here now"
        doc.save(str(path))
        result = get_document_properties(str(path))
        assert "table_word_count" in result
        assert result["table_word_count"] == 4  # "Table words here now"

    def test_table_word_count_excludes_body(self, tmp_path):
        """table_word_count only counts table text, not body."""
        path = tmp_path / "test.docx"
        doc = Document()
        doc.add_paragraph("Five words in the body")
        table = doc.add_table(rows=1, cols=1)
        table.cell(0, 0).text = "Two words"
        doc.save(str(path))
        result = get_document_properties(str(path))
        assert result["word_count"] == 5
        assert result["table_word_count"] == 2

    def test_no_tables_zero_count(self, make_docx):
        """Document with no tables has table_word_count=0."""
        path = make_docx(paragraphs=["Just body text"])
        result = get_document_properties(path)
        assert result["table_word_count"] == 0
