"""Tests for insert_line_or_paragraph_near_text copy_style_from_index (Shortcoming 2)."""
import pytest
from docx import Document
from docx.shared import Pt

from word_document_server.utils.document_utils import insert_line_or_paragraph_near_text


class TestCopyFormatting:
    """copy_style_from_index copies run-level formatting from source paragraph."""

    def test_copies_bold_and_font(self, make_docx):
        path = make_docx(paragraphs=[
            {"runs": [{"text": "Source text", "bold": True, "font_size": 14, "font_name": "Arial"}]},
            "Target paragraph",
        ])
        result = insert_line_or_paragraph_near_text(
            path,
            target_paragraph_index=0,
            line_text="New paragraph",
            position="after",
            copy_style_from_index=0,
        )
        assert "failed" not in result.lower()
        doc = Document(path)
        # The inserted paragraph should be at index 1
        new_para = doc.paragraphs[1]
        assert new_para.text == "New paragraph"
        assert len(new_para.runs) >= 1
        run = new_para.runs[0]
        assert run.bold is True
        assert run.font.name == "Arial"
        assert run.font.size == Pt(14)


class TestDefaultBehaviorUnchanged:
    """Without copy_style_from_index, behavior is unchanged (regression)."""

    def test_no_copy_style(self, make_docx):
        path = make_docx(paragraphs=[
            {"runs": [{"text": "Source text", "bold": True, "font_size": 14}]},
            "Target paragraph",
        ])
        result = insert_line_or_paragraph_near_text(
            path,
            target_paragraph_index=0,
            line_text="New paragraph",
            position="after",
        )
        assert "failed" not in result.lower()
        doc = Document(path)
        new_para = doc.paragraphs[1]
        assert new_para.text == "New paragraph"
