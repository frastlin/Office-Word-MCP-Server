"""Tests for replace_paragraph_range tool (Shortcoming 3)."""
import pytest
from docx import Document

from word_document_server.utils.document_utils import replace_paragraph_range


class TestSameCountReplacement:
    def test_replace_range_same_count(self, make_docx):
        path = make_docx(paragraphs=["A", "B", "C", "D", "E"])
        result = replace_paragraph_range(path, 1, 3, ["X", "Y", "Z"])
        assert "replaced" in result.lower()
        doc = Document(path)
        texts = [p.text for p in doc.paragraphs]
        assert texts == ["A", "X", "Y", "Z", "E"]


class TestFewerReplacements:
    def test_replace_range_fewer(self, make_docx):
        path = make_docx(paragraphs=["A", "B", "C", "D", "E"])
        result = replace_paragraph_range(path, 1, 3, ["X"])
        assert "replaced" in result.lower()
        doc = Document(path)
        texts = [p.text for p in doc.paragraphs]
        assert texts == ["A", "X", "E"]


class TestMoreReplacements:
    def test_replace_range_more(self, make_docx):
        path = make_docx(paragraphs=["A", "B", "C", "D", "E"])
        result = replace_paragraph_range(path, 1, 2, ["X", "Y", "Z", "W"])
        assert "replaced" in result.lower()
        doc = Document(path)
        texts = [p.text for p in doc.paragraphs]
        assert texts == ["A", "X", "Y", "Z", "W", "D", "E"]


class TestInvalidRange:
    def test_invalid_range_start_too_large(self, make_docx):
        path = make_docx(paragraphs=["A", "B"])
        result = replace_paragraph_range(path, 3, 5, ["X"])
        assert "invalid" in result.lower()

    def test_invalid_range_start_greater_than_end(self, make_docx):
        path = make_docx(paragraphs=["A", "B", "C"])
        result = replace_paragraph_range(path, 2, 1, ["X"])
        assert "invalid" in result.lower()


class TestStyleParameter:
    def test_style_applied(self, make_docx):
        path = make_docx(paragraphs=["A", "B", "C"])
        replace_paragraph_range(path, 0, 1, ["X", "Y"], style="Heading 1")
        doc = Document(path)
        assert doc.paragraphs[0].style.name == "Heading 1"
        assert doc.paragraphs[1].style.name == "Heading 1"


class TestSingleParagraphRange:
    def test_single_paragraph(self, make_docx):
        path = make_docx(paragraphs=["A", "B", "C"])
        replace_paragraph_range(path, 1, 1, ["X"])
        doc = Document(path)
        texts = [p.text for p in doc.paragraphs]
        assert texts == ["A", "X", "C"]
